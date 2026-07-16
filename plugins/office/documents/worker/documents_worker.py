#!/usr/bin/env python3
"""One-shot Documents Managed Worker for SheJane Plugin Action Protocol v1."""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections.abc import Iterable
from pathlib import Path, PurePosixPath
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PLUGIN_ID = "org.shejane.documents"
RUNTIME_ASSET_ID = "org.libreoffice.runtime"
ACTION_IDS = {
    "document.read",
    "document.create",
    "document.edit",
    "document.render",
}
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MARKDOWN_LIMIT = 60_000


class DocumentActionError(RuntimeError):
    def __init__(self, code: str, message: str, *, retryable: bool = False) -> None:
        super().__init__(message)
        self.code = code
        self.retryable = retryable


def _reply(request_id: int, result: object) -> None:
    print(
        json.dumps(
            {"jsonrpc": "2.0", "id": request_id, "result": result},
            ensure_ascii=False,
            separators=(",", ":"),
        ),
        flush=True,
    )


def _request(expected_method: str) -> dict[str, Any]:
    try:
        payload = json.loads(sys.stdin.readline())
    except (EOFError, json.JSONDecodeError) as exc:
        raise DocumentActionError("protocol_violation", "invalid protocol frame") from exc
    if (
        not isinstance(payload, dict)
        or payload.get("jsonrpc") != "2.0"
        or payload.get("method") != expected_method
        or not isinstance(payload.get("params"), dict)
    ):
        raise DocumentActionError("protocol_violation", f"expected {expected_method}")
    return payload


def _output_root() -> Path:
    try:
        root = Path(os.environ["SHEJANE_PLUGIN_OUTPUT_ROOT"]).resolve(strict=True)
    except (KeyError, OSError) as exc:
        raise DocumentActionError("invalid_invocation", "output staging is unavailable") from exc
    return root


def _materialized_input(invocation: dict[str, Any], input_id: str) -> Path:
    references = invocation.get("inputs")
    if not isinstance(references, list):
        raise DocumentActionError("invalid_invocation", "document input is unavailable")
    reference = next(
        (item for item in references if isinstance(item, dict) and item.get("id") == input_id),
        None,
    )
    if reference is None or reference.get("media_type") != DOCX_MIME:
        raise DocumentActionError("invalid_invocation", "a DOCX input is required")
    try:
        relative = PurePosixPath(str(reference["path"])).relative_to("/input")
    except (KeyError, ValueError) as exc:
        raise DocumentActionError("invalid_invocation", "document input path is invalid") from exc
    if not relative.parts or any(part in {"", ".", ".."} for part in relative.parts):
        raise DocumentActionError("invalid_invocation", "document input path is invalid")
    try:
        root = Path(os.environ["SHEJANE_PLUGIN_INPUT_ROOT"]).resolve(strict=True)
        source = root.joinpath(*relative.parts)
        if source.is_symlink() or not source.is_file():
            raise OSError
        source = source.resolve(strict=True)
        source.relative_to(root)
        data = source.read_bytes()
    except (KeyError, OSError, ValueError) as exc:
        raise DocumentActionError("invalid_invocation", "document input is unavailable") from exc
    if len(data) != reference.get("size_bytes") or hashlib.sha256(
        data
    ).hexdigest() != reference.get("sha256"):
        raise DocumentActionError("invalid_invocation", "document input digest changed")
    return source


def _load_docx(path: Path) -> DocumentType:
    try:
        prefix = path.read_bytes()[:8]
    except OSError as exc:
        raise DocumentActionError("invalid_invocation", "document input is unavailable") from exc
    if prefix.startswith(bytes.fromhex("d0cf11e0")):
        raise DocumentActionError(
            "document_encrypted_or_legacy",
            "encrypted or legacy binary Word documents are unsupported",
        )
    if not prefix.startswith(b"PK"):
        raise DocumentActionError("document_corrupt", "DOCX package is corrupt")
    try:
        return Document(str(path))
    except Exception as exc:
        raise DocumentActionError("document_corrupt", "DOCX package is corrupt") from exc


def _heading_level(paragraph: Paragraph) -> int | None:
    name = str(getattr(paragraph.style, "name", "") or "")
    match = re.fullmatch(r"Heading ([1-9])", name, flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def _is_list(paragraph: Paragraph) -> bool:
    properties = paragraph._p.pPr
    style = str(getattr(paragraph.style, "name", "") or "")
    return style.startswith("List ") or (properties is not None and properties.numPr is not None)


def _table_markdown(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    escaped = [[cell.replace("|", "\\|").replace("\n", " ") for cell in row] for row in normalized]
    return [
        "| " + " | ".join(escaped[0]) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
        *("| " + " | ".join(row) + " |" for row in escaped[1:]),
    ]


def _document_read(invocation: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, str]]]:
    arguments = invocation["arguments"]
    source = _materialized_input(invocation, str(arguments["input_id"]))
    document = _load_docx(source)
    headings: list[dict[str, Any]] = []
    markdown: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        level = _heading_level(paragraph)
        if level is not None:
            headings.append({"level": level, "text": text})
            markdown.append(f"{'#' * level} {text}")
        elif _is_list(paragraph):
            markdown.append(f"- {text}")
        else:
            markdown.append(text)
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        markdown.append("\n".join(_table_markdown(rows)))

    headers = [
        "\n".join(paragraph.text for paragraph in section.header.paragraphs).strip()
        for section in document.sections
    ]
    footers = [
        "\n".join(paragraph.text for paragraph in section.footer.paragraphs).strip()
        for section in document.sections
    ]
    comments = list(getattr(document, "comments", ()))
    rendered = "\n\n".join(markdown) if arguments.get("include_markdown", True) else ""
    truncated = len(rendered) > MARKDOWN_LIMIT
    rendered = rendered[:MARKDOWN_LIMIT]
    warnings = ["markdown truncated at 60000 characters"] if truncated else []
    return (
        {
            "title": str(document.core_properties.title or ""),
            "author": str(document.core_properties.author or ""),
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "comment_count": len(comments),
            "headings": headings,
            "headers": headers,
            "footers": footers,
            "markdown": rendered,
            "truncated": truncated,
            "warnings": warnings,
        },
        [],
    )


def _safe_output_name(value: str) -> str:
    if (
        not value.lower().endswith(".docx")
        or value in {".docx", "..docx"}
        or "/" in value
        or "\\" in value
        or "\x00" in value
    ):
        raise DocumentActionError("invalid_invocation", "output filename is invalid")
    return value


def _atomic_save(document: DocumentType, name: str) -> Path:
    output = _output_root()
    target = output / _safe_output_name(name)
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=".document-",
        suffix=".docx",
        dir=output,
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        document.save(str(temporary))
        _normalize_ooxml_archive(temporary)
        _load_docx(temporary)
        os.replace(temporary, target)
    except DocumentActionError:
        raise
    except Exception as exc:
        raise DocumentActionError(
            "document_write_failed", "DOCX output could not be written"
        ) from exc
    finally:
        temporary.unlink(missing_ok=True)
    return target


def _normalize_ooxml_archive(path: Path) -> None:
    normalized = path.with_suffix(".normalized.docx")
    try:
        with (
            zipfile.ZipFile(path, "r") as source,
            zipfile.ZipFile(normalized, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as target,
        ):
            for original in sorted(source.infolist(), key=lambda item: item.filename):
                if original.is_dir():
                    continue
                info = zipfile.ZipInfo(original.filename, date_time=(1980, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.create_system = 3
                info.external_attr = 0o600 << 16
                target.writestr(info, source.read(original))
        os.replace(normalized, path)
    except (OSError, zipfile.BadZipFile) as exc:
        raise DocumentActionError(
            "document_write_failed", "DOCX output could not be normalized"
        ) from exc
    finally:
        normalized.unlink(missing_ok=True)


def _set_core_properties(document: DocumentType, metadata: dict[str, Any]) -> None:
    for key in ("title", "author", "subject", "keywords"):
        if key in metadata:
            setattr(document.core_properties, key, str(metadata[key]))


def _document_create(invocation: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, str]]]:
    arguments = invocation["arguments"]
    document = Document()
    _set_core_properties(document, dict(arguments.get("metadata") or {}))
    if "header" in arguments:
        document.sections[0].header.paragraphs[0].text = str(arguments["header"])
    if "footer" in arguments:
        document.sections[0].footer.paragraphs[0].text = str(arguments["footer"])
    for block in arguments["blocks"]:
        block_type = block["type"]
        if block_type == "paragraph":
            try:
                document.add_paragraph(str(block["text"]), style=block.get("style"))
            except KeyError as exc:
                raise DocumentActionError(
                    "invalid_invocation", "paragraph style is unavailable"
                ) from exc
        elif block_type == "heading":
            document.add_heading(str(block["text"]), level=int(block["level"]))
        elif block_type in {"bullet_list", "numbered_list"}:
            style = "List Bullet" if block_type == "bullet_list" else "List Number"
            for item in block["items"]:
                document.add_paragraph(str(item), style=style)
        elif block_type == "table":
            rows = block["rows"]
            width = len(rows[0])
            if any(len(row) != width for row in rows):
                raise DocumentActionError("invalid_invocation", "table rows must have equal width")
            table = document.add_table(rows=len(rows), cols=width)
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    table.cell(row_index, column_index).text = str(value)
        elif block_type == "page_break":
            document.add_page_break()
        else:
            raise DocumentActionError("invalid_invocation", "document block type is unsupported")

    filename = _safe_output_name(str(arguments["filename"]))
    target = _atomic_save(document, filename)
    return (
        {
            "filename": filename,
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "warnings": [],
        },
        [{"path": f"/output/{target.name}", "media_type": DOCX_MIME, "name": target.name}],
    )


def _iter_paragraphs(document: DocumentType) -> Iterable[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _find_paragraph(document: DocumentType, target: str) -> Paragraph | None:
    return next(
        (paragraph for paragraph in _iter_paragraphs(document) if paragraph.text == target), None
    )


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _insert_after(paragraph: Paragraph, text: str, style: str | None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style:
        try:
            inserted.style = style
        except KeyError as exc:
            raise DocumentActionError(
                "invalid_invocation", "paragraph style is unavailable"
            ) from exc
    inserted.add_run(text)
    return inserted


def _document_edit(invocation: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, str]]]:
    arguments = invocation["arguments"]
    source = _materialized_input(invocation, str(arguments["input_id"]))
    document = _load_docx(source)
    changes: list[str] = []
    for operation in arguments["operations"]:
        operation_type = operation["type"]
        if operation_type == "find_replace":
            remaining = int(operation.get("max_replacements", 10_000))
            count = 0
            for paragraph in _iter_paragraphs(document):
                if remaining <= 0:
                    break
                occurrences = min(paragraph.text.count(operation["find"]), remaining)
                if occurrences:
                    _set_paragraph_text(
                        paragraph,
                        paragraph.text.replace(
                            operation["find"], operation["replace"], occurrences
                        ),
                    )
                    count += occurrences
                    remaining -= occurrences
            if count == 0:
                raise DocumentActionError("document_target_not_found", "find text was not found")
            changes.append(f"find_replace:{count}")
        elif operation_type == "insert_paragraph":
            if operation.get("after"):
                anchor = _find_paragraph(document, str(operation["after"]))
                if anchor is None:
                    raise DocumentActionError(
                        "document_target_not_found", "paragraph anchor was not found"
                    )
                _insert_after(anchor, str(operation["text"]), operation.get("style"))
            else:
                try:
                    document.add_paragraph(str(operation["text"]), style=operation.get("style"))
                except KeyError as exc:
                    raise DocumentActionError(
                        "invalid_invocation", "paragraph style is unavailable"
                    ) from exc
            changes.append("insert_paragraph:1")
        elif operation_type in {"update_paragraph", "delete_paragraph", "apply_style"}:
            paragraph = _find_paragraph(document, str(operation["target"]))
            if paragraph is None:
                raise DocumentActionError(
                    "document_target_not_found", "paragraph target was not found"
                )
            if operation_type == "update_paragraph":
                _set_paragraph_text(paragraph, str(operation["text"]))
            elif operation_type == "delete_paragraph":
                parent = paragraph._p.getparent()
                parent.remove(paragraph._p)
            else:
                try:
                    paragraph.style = str(operation["style"])
                except KeyError as exc:
                    raise DocumentActionError(
                        "invalid_invocation", "paragraph style is unavailable"
                    ) from exc
            changes.append(f"{operation_type}:1")
        elif operation_type in {"set_header", "set_footer"}:
            sections = list(document.sections)
            indexes = (
                [int(operation["section"])] if "section" in operation else range(len(sections))
            )
            try:
                for index in indexes:
                    container = (
                        sections[index].header
                        if operation_type == "set_header"
                        else sections[index].footer
                    )
                    container.paragraphs[0].text = str(operation["text"])
            except IndexError as exc:
                raise DocumentActionError(
                    "invalid_invocation", "document section is unavailable"
                ) from exc
            changes.append(f"{operation_type}:1")
        elif operation_type == "set_metadata":
            _set_core_properties(document, {str(operation["property"]): str(operation["value"])})
            changes.append("set_metadata:1")
        elif operation_type == "add_page_break":
            if operation.get("after"):
                paragraph = _find_paragraph(document, str(operation["after"]))
                if paragraph is None:
                    raise DocumentActionError(
                        "document_target_not_found", "paragraph anchor was not found"
                    )
                paragraph.add_run().add_break(WD_BREAK.PAGE)
            else:
                document.add_page_break()
            changes.append("add_page_break:1")
        else:
            raise DocumentActionError(
                "invalid_invocation", "document edit operation is unsupported"
            )

    default_name = f"{source.stem}.edited.docx"
    filename = _safe_output_name(str(arguments.get("output_filename") or default_name))
    target = _atomic_save(document, filename)
    return (
        {
            "filename": filename,
            "operation_count": len(arguments["operations"]),
            "changes": changes,
            "warnings": [],
        },
        [{"path": f"/output/{target.name}", "media_type": DOCX_MIME, "name": target.name}],
    )


def _runtime_tool_paths() -> tuple[Path, Path]:
    try:
        mapping = json.loads(os.environ["SHEJANE_PLUGIN_RUNTIME_ASSETS"])
        root = Path(mapping[RUNTIME_ASSET_ID]).resolve(strict=True)
        config = json.loads((root / "office-runtime.json").read_text(encoding="utf-8"))
    except (KeyError, OSError, ValueError, json.JSONDecodeError) as exc:
        raise DocumentActionError(
            "engine_unavailable", "Office runtime asset is unavailable"
        ) from exc
    tools: list[Path] = []
    for key in ("soffice", "mutool"):
        raw = str(config.get(key, ""))
        relative = PurePosixPath(raw)
        if (
            not raw
            or relative.is_absolute()
            or any(part in {"", ".", ".."} for part in relative.parts)
        ):
            raise DocumentActionError("engine_unavailable", "Office runtime asset is invalid")
        candidate = root.joinpath(*relative.parts)
        try:
            candidate = candidate.resolve(strict=True)
            candidate.relative_to(root)
        except (OSError, ValueError) as exc:
            raise DocumentActionError(
                "engine_unavailable", "Office runtime asset is invalid"
            ) from exc
        if not candidate.is_file() or not os.access(candidate, os.X_OK):
            raise DocumentActionError(
                "engine_unavailable", "Office runtime executable is unavailable"
            )
        tools.append(candidate)
    return tools[0], tools[1]


def _run_tool(command: list[str], *, cwd: Path, timeout: float, env: dict[str, str]) -> str:
    try:
        completed = subprocess.run(
            command,
            cwd=cwd,
            env=env,
            stdin=subprocess.DEVNULL,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            timeout=timeout,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        raise DocumentActionError("render_failed", "document renderer did not complete") from exc
    if completed.returncode != 0:
        raise DocumentActionError("render_failed", "document renderer rejected the input")
    return completed.stdout[-16_384:].decode("utf-8", errors="replace")


def _document_render(invocation: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, str]]]:
    arguments = invocation["arguments"]
    source = _materialized_input(invocation, str(arguments["input_id"]))
    _load_docx(source)
    soffice, mutool = _runtime_tool_paths()
    output = _output_root()
    profile = output / ".libreoffice-profile"
    temporary = output / ".runtime-tmp"
    profile.mkdir(mode=0o700, exist_ok=False)
    temporary.mkdir(mode=0o700, exist_ok=False)
    timeout = max(1.0, min(300.0, float(invocation["limits"]["timeout_ms"]) / 1000 * 0.8))
    env = {
        "HOME": str(profile),
        "TMPDIR": str(temporary),
        "TEMP": str(temporary),
        "TMP": str(temporary),
        "LANG": "C.UTF-8",
        "LC_ALL": "C.UTF-8",
        "TZ": "UTC",
        "PATH": os.pathsep.join(sorted({str(soffice.parent), str(mutool.parent)})),
    }
    try:
        _run_tool(
            [
                str(soffice),
                "--headless",
                "--nologo",
                "--nodefault",
                "--nolockcheck",
                "--norestore",
                f"-env:UserInstallation={profile.as_uri()}",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(output),
                str(source),
            ],
            cwd=output,
            timeout=timeout,
            env=env,
        )
        pdf = output / f"{source.stem}.pdf"
        if not pdf.is_file() or not pdf.read_bytes().startswith(b"%PDF"):
            raise DocumentActionError("render_failed", "document renderer produced no valid PDF")
        info = _run_tool([str(mutool), "info", str(pdf)], cwd=output, timeout=timeout, env=env)
        page_match = re.search(r"(?:Pages|pages)\s*:\s*(\d+)", info)
        if page_match is None or int(page_match.group(1)) < 1:
            raise DocumentActionError("render_failed", "PDF page count is unavailable")
        page_count = int(page_match.group(1))
        include_png = bool(arguments.get("include_png", True))
        max_pages = int(arguments.get("max_pages", 20))
        rendered_pages = min(page_count, max_pages) if include_png else 0
        png_names: list[str] = []
        if rendered_pages:
            pattern = output / f"{source.stem}.page-%04d.png"
            _run_tool(
                [
                    str(mutool),
                    "draw",
                    "-q",
                    "-r",
                    str(int(arguments.get("dpi", 144))),
                    "-o",
                    str(pattern),
                    str(pdf),
                    f"1-{rendered_pages}",
                ],
                cwd=output,
                timeout=timeout,
                env=env,
            )
            for page in range(1, rendered_pages + 1):
                name = f"{source.stem}.page-{page:04d}.png"
                candidate = output / name
                if not candidate.is_file():
                    raise DocumentActionError("render_failed", "PNG page rendering was incomplete")
                png_names.append(name)
        warnings = (
            ["PNG preview limited by max_pages"]
            if rendered_pages < page_count and include_png
            else []
        )
        artifacts = [
            {"path": f"/output/{pdf.name}", "media_type": "application/pdf", "name": pdf.name},
            *(
                {"path": f"/output/{name}", "media_type": "image/png", "name": name}
                for name in png_names
            ),
        ]
        return (
            {
                "page_count": page_count,
                "rendered_pages": rendered_pages,
                "pdf_name": pdf.name,
                "png_names": png_names,
                "warnings": warnings,
            },
            artifacts,
        )
    finally:
        shutil.rmtree(profile, ignore_errors=True)
        shutil.rmtree(temporary, ignore_errors=True)


def _success(
    invocation: dict[str, Any],
    output: dict[str, Any],
    artifacts: list[dict[str, str]],
) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "invocation_id": invocation["invocation_id"],
        "operation_id": invocation["operation_id"],
        "status": "succeeded",
        "output": output,
        "artifacts": artifacts,
    }


def _failure(invocation: dict[str, Any], error: DocumentActionError) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "invocation_id": invocation["invocation_id"],
        "operation_id": invocation["operation_id"],
        "status": "failed",
        "error": {
            "code": error.code,
            "message": str(error),
            "retryable": error.retryable,
        },
        "artifacts": [],
    }


def _invoke(invocation: dict[str, Any]) -> dict[str, Any]:
    action = invocation.get("action")
    if not isinstance(action, dict) or action.get("plugin_id") != PLUGIN_ID:
        raise DocumentActionError("invalid_invocation", "plugin identity does not match")
    action_id = str(action.get("action_id") or "")
    handlers = {
        "document.read": _document_read,
        "document.create": _document_create,
        "document.edit": _document_edit,
        "document.render": _document_render,
    }
    handler = handlers.get(action_id)
    if handler is None:
        raise DocumentActionError("invalid_invocation", "document Action is unsupported")
    output, artifacts = handler(invocation)
    return _success(invocation, output, artifacts)


def main() -> None:
    initialize = _request("initialize")
    params = initialize["params"]
    if (
        params.get("protocol_version") != 1
        or params.get("plugin_id") != PLUGIN_ID
        or not set(params.get("actions", ())).issubset(ACTION_IDS)
    ):
        raise DocumentActionError("incompatible_runtime", "worker initialization is incompatible")
    _reply(
        int(initialize["id"]),
        {
            "protocol_version": 1,
            "process_isolated": True,
            "access_isolated": os.environ.get("SHEJANE_PLUGIN_ACCESS_ISOLATED") == "1",
            "resource_isolated": os.environ.get("SHEJANE_PLUGIN_RESOURCE_ISOLATED") == "1",
            "sandboxed": os.environ.get("SHEJANE_PLUGIN_SANDBOXED") == "1",
        },
    )
    invoke = _request("invoke")
    invocation = invoke["params"]
    try:
        result = _invoke(invocation)
    except DocumentActionError as exc:
        result = _failure(invocation, exc)
    except Exception as exc:
        if os.environ.get("SHEJANE_PLUGIN_DEBUG") == "1":
            print(f"{type(exc).__name__}: {exc}", file=sys.stderr, flush=True)
        result = _failure(
            invocation,
            DocumentActionError("plugin_failed", "document worker failed"),
        )
    _reply(int(invoke["id"]), result)
    shutdown = _request("shutdown")
    _reply(int(shutdown["id"]), {})


if __name__ == "__main__":
    main()
