"""Unit tests for office.read and office.outline.

The fixtures here are built on-the-fly with python-docx / openpyxl so we
don't have to check binary files into the repo. Each test owns its own
tmp_path-scoped file.

Office tools don't proxy through any external service (markitdown +
python-docx + openpyxl all run in-process), so these are vanilla unit
tests — no MockTransport, no settings override needed.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook

from local_host.tools.office import (
    OFFICE_READ_TOOLS,
    OFFICE_WRITE_TOOLS,
    edited_copy_path,
    office_add_row,
    office_apply_style,
    office_delete_paragraph,
    office_find_replace,
    office_insert_paragraph,
    office_merge_cells,
    office_outline,
    office_read,
    office_read_range,
    office_set_cell_format,
    office_set_cells,
    office_set_formula,
    office_update_paragraph,
)


def _make_docx(tmp_path: Path) -> Path:
    """Build a tiny .docx with heading + paragraphs + a table."""
    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("First paragraph with some prose.")
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph("Methodology details here.")
    doc.add_paragraph("Another methodology paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Year"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "2025"
    table.rows[1].cells[1].text = "42"
    doc.save(p)
    return p


def _make_xlsx(tmp_path: Path) -> Path:
    """Build a tiny .xlsx with two sheets."""
    p = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sales"
    ws1.append(["Quarter", "Revenue"])
    ws1.append(["Q1", 100])
    ws1.append(["Q2", 150])
    ws2 = wb.create_sheet("Costs")
    ws2.append(["Item", "Amount"])
    ws2.append(["Rent", 50])
    wb.save(p)
    return p


def test_office_read_docx_returns_markdown(tmp_path: Path) -> None:
    """office.read on a .docx → markdown with the headings preserved."""
    path = _make_docx(tmp_path)
    result = office_read.func(path=str(path))
    assert result["ok"] == "true"
    assert result["kind"] == "word"
    md = result["markdown"]
    # Heading text should survive in the markdown output. markitdown's
    # exact heading-level mapping varies; assert the text is there.
    assert "Executive Summary" in md
    assert "Methodology" in md
    # Table contents should also surface.
    assert "Year" in md and "Value" in md
    # Not truncated for a tiny doc.
    assert result["truncated"] == "false"


def test_office_read_xlsx_returns_markdown(tmp_path: Path) -> None:
    """office.read on a .xlsx → markdown with both sheets."""
    path = _make_xlsx(tmp_path)
    result = office_read.func(path=str(path))
    assert result["ok"] == "true"
    assert result["kind"] == "excel"
    md = result["markdown"]
    # Both sheet names should appear (markitdown sections by sheet).
    assert "Sales" in md
    assert "Costs" in md
    # Cell contents should surface.
    assert "Q1" in md and "100" in md


def test_office_read_missing_file_returns_error(tmp_path: Path) -> None:
    """office.read on a non-existent path → ok=false, no crash."""
    result = office_read.func(path=str(tmp_path / "missing.docx"))
    assert result["ok"] == "false"
    assert "not found" in result["error"]


def test_office_read_wrong_extension_returns_error(tmp_path: Path) -> None:
    """office.read on a .txt → ok=false steering the LLM to read_file."""
    txt = tmp_path / "note.txt"
    txt.write_text("plain text")
    result = office_read.func(path=str(txt))
    assert result["ok"] == "false"
    # Error message should mention the supported extensions so the LLM
    # can self-correct on the next call.
    assert ".docx" in result["error"] and ".xlsx" in result["error"]


def test_office_read_empty_path_returns_error() -> None:
    """office.read with empty path → ok=false (matches workspace.open's behavior)."""
    result = office_read.func(path="")
    assert result["ok"] == "false"
    assert result["error"] == "path required"


def test_office_outline_docx_lists_headings(tmp_path: Path) -> None:
    """office.outline on .docx surfaces heading text + paragraph/table counts."""
    path = _make_docx(tmp_path)
    result = office_outline.func(path=str(path))
    assert result["ok"] == "true"
    assert result["kind"] == "word"
    heading_texts = [h["text"] for h in result["headings"]]
    assert "Executive Summary" in heading_texts
    assert "Methodology" in heading_texts
    # Two top-level + one subsection means at least one heading per level.
    levels = {h["level"] for h in result["headings"]}
    assert 1 in levels and 2 in levels
    # Paragraphs: 2 headings + 3 body paragraphs (heading paragraphs count
    # too in python-docx's iteration); >= 5 is the safe bound.
    assert result["paragraph_count"] >= 5
    assert result["table_count"] == 1


def test_office_outline_xlsx_lists_sheets(tmp_path: Path) -> None:
    """office.outline on .xlsx surfaces sheet names + dimensions."""
    path = _make_xlsx(tmp_path)
    result = office_outline.func(path=str(path))
    assert result["ok"] == "true"
    assert result["kind"] == "excel"
    sheet_names = [s["name"] for s in result["sheets"]]
    assert sheet_names == ["Sales", "Costs"]
    sales = next(s for s in result["sheets"] if s["name"] == "Sales")
    assert sales["rows"] == 3  # header + 2 data rows
    assert sales["columns"] == 2


def test_office_outline_wrong_extension_returns_error(tmp_path: Path) -> None:
    """office.outline rejects non-office files just like office.read."""
    txt = tmp_path / "x.csv"
    txt.write_text("a,b\n1,2\n")
    result = office_outline.func(path=str(txt))
    assert result["ok"] == "false"
    assert "unsupported extension" in result["error"]


def test_office_tools_registered_under_canonical_names() -> None:
    """Sanity: OFFICE_READ_TOOLS exposes the read tools by their dotted names."""
    names = sorted(t.name for t in OFFICE_READ_TOOLS)
    assert names == ["office.outline", "office.read", "office.read_range"]


def test_office_read_truncates_large_output(tmp_path: Path) -> None:
    """office.read above the 60_000 char cap → truncated="true" + the cap."""
    from local_host.tools import office as office_module

    # Build a .docx with way more than 60k chars of body text by repeating
    # a paragraph many times. Cheaper than the LLM cap test using fixtures.
    p = tmp_path / "big.docx"
    doc = Document()
    para = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 50  # ~2900 chars/para
    for _ in range(25):  # ~72500 chars total before markdown overhead
        doc.add_paragraph(para)
    doc.save(p)
    result = office_read.func(path=str(p))
    assert result["ok"] == "true"
    assert result["truncated"] == "true"
    # The cap is enforced on the returned markdown; allow some slack for
    # the trailing "…(truncated, full size = N chars)" suffix.
    assert len(result["markdown"]) >= office_module._MARKDOWN_CHAR_CAP
    assert len(result["markdown"]) < office_module._MARKDOWN_CHAR_CAP + 200


@pytest.mark.parametrize("ext", [".doc", ".xls", ".pptx", ".pdf"])
def test_office_read_rejects_non_supported_office_extensions(tmp_path: Path, ext: str) -> None:
    """office.read explicitly rejects .doc / .xls / .pptx / .pdf — we only
    support OOXML formats in Phase 1. Phase 2 may revisit .doc/.xls (the
    legacy binary formats require a separate parser)."""
    p = tmp_path / f"x{ext}"
    p.write_bytes(b"\x00")  # dummy bytes, validation happens on extension first
    result = office_read.func(path=str(p))
    assert result["ok"] == "false"
    assert ext in result["error"] or "unsupported extension" in result["error"]


# ═══════════════════════════════════════════════════════════════════════
# Phase 2: write tools (copy-on-first-write)
# ═══════════════════════════════════════════════════════════════════════


def _sha256(path: Path) -> str:
    import hashlib

    return hashlib.sha256(path.read_bytes()).hexdigest()


def _docx_paragraphs(path: Path) -> list[str]:
    return [p.text for p in Document(path).paragraphs]


def _xlsx_cell(path: Path, sheet: str, cell: str):
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=False)
    try:
        return wb[sheet][cell].value
    finally:
        wb.close()


def test_write_tools_registered_under_canonical_names() -> None:
    """Sanity: OFFICE_WRITE_TOOLS exposes the 10 write tools by name."""
    names = sorted(t.name for t in OFFICE_WRITE_TOOLS)
    assert names == sorted(
        [
            "office.find_replace",
            "office.insert_paragraph",
            "office.update_paragraph",
            "office.delete_paragraph",
            "office.apply_style",
            "office.set_cells",
            "office.set_formula",
            "office.set_cell_format",
            "office.merge_cells",
            "office.add_row",
        ]
    )


def test_edited_copy_path_helper() -> None:
    """`<basename>.edited.<ext>` lives in the same dir; already-edited
    paths are returned unchanged."""
    assert edited_copy_path("/x/y/report.docx") == "/x/y/report.edited.docx"
    assert edited_copy_path("/x/y/report.edited.docx") == "/x/y/report.edited.docx"
    assert edited_copy_path("/x/y/q4.xlsx") == "/x/y/q4.edited.xlsx"


def test_find_replace_writes_to_copy_and_preserves_original(tmp_path: Path) -> None:
    """Core contract: original SHA256 unchanged; .edited contains the replacement."""
    p = _make_docx(tmp_path)
    sha_before = _sha256(p)
    result = office_find_replace.func(
        path=str(p),
        find="First paragraph",
        replace="Updated first sentence",
    )
    assert result["ok"] == "true", result
    edited = Path(result["edited_path"])
    assert edited.exists()
    assert edited != p
    # Original byte-for-byte unchanged.
    assert _sha256(p) == sha_before
    # Edited copy has the new text.
    paragraphs = _docx_paragraphs(edited)
    assert any("Updated first sentence" in para for para in paragraphs)
    assert all("First paragraph with some prose." not in para for para in paragraphs)
    assert result["summary"]["replaced"] == 1


def test_find_replace_respects_count_limit(tmp_path: Path) -> None:
    """count=N caps the global number of replacements."""
    p = tmp_path / "repeats.docx"
    doc = Document()
    for _ in range(5):
        doc.add_paragraph("apple banana apple")
    doc.save(p)
    result = office_find_replace.func(path=str(p), find="apple", replace="X", count=3)
    assert result["ok"] == "true"
    assert result["summary"]["replaced"] == 3
    text = "\n".join(_docx_paragraphs(Path(result["edited_path"])))
    # Started with 5 paragraphs × 2 = 10 "apple"; 3 replaced → 7 remain.
    assert text.count("apple") == 7
    assert text.count("X") == 3


def test_repeat_write_targets_same_edited_file(tmp_path: Path) -> None:
    """Two writes against the original land in ONE `.edited.docx`, the
    second building on the first."""
    p = _make_docx(tmp_path)
    r1 = office_find_replace.func(path=str(p), find="Methodology", replace="Approach")
    assert r1["ok"] == "true"
    edited = Path(r1["edited_path"])
    sha_after_first = _sha256(edited)

    r2 = office_find_replace.func(path=str(p), find="Executive", replace="Top-line")
    assert r2["ok"] == "true"
    # Same edited path both times.
    assert r2["edited_path"] == r1["edited_path"]
    # Edited file changed between writes (we built on top of the first).
    assert _sha256(edited) != sha_after_first
    # Second edit incorporates BOTH changes.
    texts = "\n".join(_docx_paragraphs(edited))
    assert "Approach" in texts
    assert "Top-line" in texts
    assert "Methodology" not in texts
    assert "Executive" not in texts


def test_insert_paragraph_after_anchor(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    result = office_insert_paragraph.func(
        path=str(p),
        anchor="Methodology",
        content="Inserted sentence.",
        position="after",
        style=None,
    )
    assert result["ok"] == "true", result
    edited_paragraphs = _docx_paragraphs(Path(result["edited_path"]))
    # The inserted paragraph appears AFTER the Methodology heading.
    idx_methodology = next(i for i, t in enumerate(edited_paragraphs) if "Methodology" in t)
    assert edited_paragraphs[idx_methodology + 1] == "Inserted sentence."


def test_insert_paragraph_at_end_with_style(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    before_count = len(_docx_paragraphs(p))
    result = office_insert_paragraph.func(
        path=str(p),
        anchor="__END__",
        content="Conclusion",
        style="Heading 1",
    )
    assert result["ok"] == "true"
    paragraphs = _docx_paragraphs(Path(result["edited_path"]))
    assert paragraphs[-1] == "Conclusion"
    assert len(paragraphs) == before_count + 1


def test_insert_paragraph_missing_anchor_returns_error(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    result = office_insert_paragraph.func(
        path=str(p),
        anchor="this-string-doesnt-exist",
        content="…",
    )
    assert result["ok"] == "false"
    assert "anchor not found" in result["error"]


def test_update_paragraph_replaces_full_text(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    result = office_update_paragraph.func(
        path=str(p),
        target="First paragraph",
        content="Entirely rewritten paragraph body.",
    )
    assert result["ok"] == "true"
    paragraphs = _docx_paragraphs(Path(result["edited_path"]))
    assert any("Entirely rewritten paragraph body." in t for t in paragraphs)
    assert all("First paragraph with some prose." not in t for t in paragraphs)


def test_delete_paragraph_removes_one(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    before = _docx_paragraphs(p)
    result = office_delete_paragraph.func(path=str(p), target="Methodology details")
    assert result["ok"] == "true"
    after = _docx_paragraphs(Path(result["edited_path"]))
    assert len(after) == len(before) - 1
    assert not any("Methodology details" in t for t in after)


def test_apply_style_changes_style(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    result = office_apply_style.func(
        path=str(p),
        target="First paragraph",
        style="Heading 2",
    )
    assert result["ok"] == "true"
    assert result["summary"]["new_style"] == "Heading 2"
    # Reopen and verify the style stuck.
    doc = Document(Path(result["edited_path"]))
    match = next(p for p in doc.paragraphs if "First paragraph" in p.text)
    assert match.style.name == "Heading 2"


def test_set_cells_writes_2d_block(tmp_path: Path) -> None:
    p = _make_xlsx(tmp_path)
    sha_before = _sha256(p)
    result = office_set_cells.func(
        path=str(p),
        sheet="Sales",
        range="C1:D2",
        values=[["X", "Y"], [10, 20]],
    )
    assert result["ok"] == "true", result
    # Original untouched.
    assert _sha256(p) == sha_before
    edited = Path(result["edited_path"])
    assert _xlsx_cell(edited, "Sales", "C1") == "X"
    assert _xlsx_cell(edited, "Sales", "D2") == 20
    # Unrelated cells unaffected.
    assert _xlsx_cell(edited, "Sales", "A1") == "Quarter"
    assert result["summary"]["cells_written"] == 4


def test_set_formula_writes_formula_literal(tmp_path: Path) -> None:
    p = _make_xlsx(tmp_path)
    result = office_set_formula.func(
        path=str(p),
        sheet="Sales",
        cell="C5",
        formula="=SUM(B2:B3)",
    )
    assert result["ok"] == "true"
    # data_only=False reads the formula text (openpyxl can't evaluate;
    # the cached value is None until Excel opens the file).
    assert _xlsx_cell(Path(result["edited_path"]), "Sales", "C5") == "=SUM(B2:B3)"


def test_set_formula_rejects_missing_leading_equals(tmp_path: Path) -> None:
    p = _make_xlsx(tmp_path)
    result = office_set_formula.func(path=str(p), sheet="Sales", cell="C1", formula="SUM(A1:A3)")
    assert result["ok"] == "false"
    assert "must start with '='" in result["error"]


def test_set_cell_format_applies_bold_and_color(tmp_path: Path) -> None:
    from openpyxl import load_workbook

    p = _make_xlsx(tmp_path)
    result = office_set_cell_format.func(
        path=str(p),
        sheet="Sales",
        range="A1:B1",
        bold=True,
        font_color="#FF5722",
        bg_color="#FFFFCC",
    )
    assert result["ok"] == "true"
    wb = load_workbook(Path(result["edited_path"]))
    try:
        ws = wb["Sales"]
        cell = ws["A1"]
        assert cell.font.bold is True
        assert "FF5722" in (cell.font.color.rgb or "").upper()
        # PatternFill stores color object; coerce to string for compare.
        fill_color = str(cell.fill.start_color.rgb or "").upper()
        assert "FFFFCC" in fill_color
    finally:
        wb.close()
    assert "bold" in result["summary"]["applied"]
    assert result["summary"]["cells_formatted"] == 2


def test_set_cell_format_rejects_invalid_color(tmp_path: Path) -> None:
    p = _make_xlsx(tmp_path)
    result = office_set_cell_format.func(
        path=str(p), sheet="Sales", range="A1", font_color="not-a-hex"
    )
    assert result["ok"] == "false"
    assert "hex" in result["error"]


def test_merge_cells_records_merge(tmp_path: Path) -> None:
    from openpyxl import load_workbook

    p = _make_xlsx(tmp_path)
    result = office_merge_cells.func(path=str(p), sheet="Sales", range="A1:B1")
    assert result["ok"] == "true"
    wb = load_workbook(Path(result["edited_path"]))
    try:
        ranges = [str(r) for r in wb["Sales"].merged_cells.ranges]
        assert "A1:B1" in ranges
    finally:
        wb.close()


def test_add_row_append_and_insert(tmp_path: Path) -> None:
    p = _make_xlsx(tmp_path)

    # Append.
    r_append = office_add_row.func(
        path=str(p), sheet="Sales", values=["Q3", 200], position="append"
    )
    assert r_append["ok"] == "true"
    edited = Path(r_append["edited_path"])
    assert _xlsx_cell(edited, "Sales", "A4") == "Q3"
    assert _xlsx_cell(edited, "Sales", "B4") == 200

    # Insert at row 2 (before existing Q1 data).
    r_insert = office_add_row.func(path=str(p), sheet="Sales", values=["Q0", 50], position=2)
    assert r_insert["ok"] == "true"
    # row 2 is now Q0, row 3 (previously Q1) shifted down
    edited = Path(r_insert["edited_path"])
    assert _xlsx_cell(edited, "Sales", "A2") == "Q0"
    assert _xlsx_cell(edited, "Sales", "A3") == "Q1"


def test_read_range_returns_values_and_formulas(tmp_path: Path) -> None:
    p = _make_xlsx(tmp_path)
    # Plant a formula first so we have something to read.
    office_set_formula.func(path=str(p), sheet="Sales", cell="C2", formula="=B2*2")
    edited = Path(p).with_name("sample.edited.xlsx")
    assert edited.exists()

    result = office_read_range.func(path=str(edited), sheet="Sales", range="A1:C3")
    assert result["ok"] == "true"
    assert result["values"][0] == ["Quarter", "Revenue", None]
    assert result["values"][1][0] == "Q1"
    assert result["values"][1][1] == 100
    # Formula text surfaced at C2.
    assert result["formulas"][1][2] == "=B2*2"
    # Non-formula cells return None in the formulas grid.
    assert result["formulas"][0][0] is None


def test_read_range_rejects_docx(tmp_path: Path) -> None:
    p = _make_docx(tmp_path)
    result = office_read_range.func(path=str(p), sheet=None, range="A1:C3")
    assert result["ok"] == "false"
    assert ".xlsx" in result["error"]


def test_write_tool_rejects_unsupported_kind(tmp_path: Path) -> None:
    """find_replace is docx-only; calling on a .xlsx returns ok=false
    BEFORE creating any copy."""
    p = _make_xlsx(tmp_path)
    result = office_find_replace.func(path=str(p), find="x", replace="y")
    assert result["ok"] == "false"
    assert "requires a .docx" in result["error"]
    # Crucially: no .edited.xlsx file was created.
    assert not (tmp_path / "sample.edited.xlsx").exists()


def test_atomic_write_keeps_target_intact_on_verification_failure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """If verification fails after the tmp write, the target keeps its
    previous (last-successful) content and no tmp file is left behind."""
    from local_host.tools import office as office_module

    p = _make_docx(tmp_path)
    # First do a real edit so .edited.docx exists with known content.
    r1 = office_find_replace.func(path=str(p), find="Methodology", replace="Approach")
    assert r1["ok"] == "true"
    target = Path(r1["edited_path"])
    sha_known_good = _sha256(target)

    # Now force _verify_file to raise the next time it runs — simulates
    # a corrupted-on-save scenario.
    def _boom(path: str, kind: str) -> None:
        raise RuntimeError("simulated corruption")

    monkeypatch.setattr(office_module, "_verify_file", _boom)

    r2 = office_find_replace.func(path=str(p), find="Executive", replace="Top-line")
    assert r2["ok"] == "false"
    # Target file is unchanged (still the last known good).
    assert _sha256(target) == sha_known_good
    # No .tmp file lingering.
    assert not (tmp_path / "sample.edited.docx.tmp").exists()
