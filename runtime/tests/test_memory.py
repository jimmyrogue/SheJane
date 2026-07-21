from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from docx import Document
from langgraph.store.memory import InMemoryStore

from shejane_runtime.auth import LOCAL_OWNER_PRINCIPAL_ID
from shejane_runtime.config import Settings, reset_settings_for_tests
from shejane_runtime.tools.memory import (
    NAMESPACE,
    extract_memory_write_facts,
    make_memory_search_tool,
    memory_namespace_for_workspace,
    memory_write,
)


async def test_open_store_persists_values(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import open_store

    reset_settings_for_tests(data_dir=tmp_path)
    store, stack = await open_store()
    try:
        await store.aput(NAMESPACE, "k1", {"text": "hello"})
        assert [item.key for item in await store.asearch(NAMESPACE)] == ["k1"]
    finally:
        await stack.aclose()


def test_resolve_memory_sources(monkeypatch, tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _resolve_memory_sources

    assert _resolve_memory_sources(Settings(SHEJANE_RUNTIME_MEMORY_PATHS="")) is None
    monkeypatch.setenv("HOME", "/home/test")
    settings = Settings(SHEJANE_RUNTIME_MEMORY_PATHS="~/a.md,/abs/b.md, /spaced/c.md ")
    assert _resolve_memory_sources(settings) == [
        "/home/test/a.md",
        "/abs/b.md",
        "/spaced/c.md",
    ]
    instruction_dir = tmp_path / "project"
    instruction_dir.mkdir()
    assert _resolve_memory_sources(Settings(SHEJANE_RUNTIME_MEMORY_PATHS=str(instruction_dir))) == [
        str(instruction_dir / "AGENTS.md")
    ]


def test_memory_instruction_sources_are_read_only_backend_routes(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    agents_md = workspace / "AGENTS.md"
    agents_md.write_text("trusted instructions", encoding="utf-8")
    external = tmp_path / "external"
    external.mkdir()
    external_agents_md = external / "AGENTS.md"
    external_agents_md.write_text("external instructions", encoding="utf-8")
    external_sibling = external / "secret.txt"
    external_sibling.write_text("SECRET_NEIGHBOR", encoding="utf-8")

    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[str(agents_md), str(external_agents_md)],
    )

    trusted = backend.read(str(agents_md.resolve()))
    assert trusted.error is None
    assert trusted.file_data == {"content": "trusted instructions", "encoding": "utf-8"}
    assert (
        backend.edit(
            str(agents_md.resolve()),
            "trusted",
            "tampered",
        ).error
        == "read-only source: edits are not allowed"
    )
    assert (
        backend.edit("/AGENTS.md", "trusted", "tampered").error
        == "read-only source: edits are not allowed"
    )
    grep_result = backend.grep("SECRET_NEIGHBOR", path="/")
    assert grep_result.error is None
    assert not any(item.get("text") == "SECRET_NEIGHBOR" for item in grep_result.matches or [])
    denied = backend.read(str(external_sibling.resolve()))
    assert denied.file_data is None
    assert denied.error is not None
    assert backend.write(str((workspace / "workspace-note.md").resolve()), "allowed").error is None


def test_attachment_backend_route_exposes_only_the_selected_file(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    source = tmp_path / "uploads"
    source.mkdir()
    attachment = source / "brief.txt"
    attachment.write_text("selected attachment", encoding="utf-8")
    (source / "secret.txt").write_text("not selected", encoding="utf-8")

    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[],
        attachment_bindings=[
            {
                "source_path": str(attachment),
                "virtual_path": "/attachments/brief.txt",
            }
        ],
    )

    selected = backend.read("/attachments/brief.txt")
    assert selected.file_data == {"content": "selected attachment", "encoding": "utf-8"}
    assert backend.read("/attachments/secret.txt").file_data is None
    assert backend.write("/attachments/brief.txt", "changed").error == (
        "read-only source: writes are not allowed"
    )


def test_pdf_attachment_is_exposed_as_model_readable_text(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    attachment = tmp_path / "rental-receipt.pdf"
    attachment.write_bytes(_minimal_pdf("Rental receipt"))

    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[],
        attachment_bindings=[
            {
                "source_path": str(attachment),
                "virtual_path": "/attachments/rental-receipt.pdf",
            }
        ],
    )

    selected = backend.read("/attachments/rental-receipt.pdf")
    assert selected.error is None
    assert selected.file_data is not None
    assert selected.file_data["encoding"] == "utf-8"
    assert "Rental receipt" in selected.file_data["content"]


def test_docx_attachment_snapshot_is_exposed_as_model_readable_text(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    original = tmp_path / "contract.docx"
    document = Document()
    document.add_heading("Runtime-owned contract", level=1)
    document.add_paragraph("The immutable attachment snapshot is readable.")
    document.save(original)
    snapshot = tmp_path / "2170c9b2f13c22ccce526b85594ec9d4"
    snapshot.write_bytes(original.read_bytes())
    original.unlink()

    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[],
        attachment_bindings=[
            {
                "source_path": str(snapshot),
                "virtual_path": "/attachments/contract.docx",
            }
        ],
    )

    selected = backend.read("/attachments/contract.docx")
    assert selected.error is None
    assert selected.file_data is not None
    assert "Runtime-owned contract" in selected.file_data["content"]
    assert "immutable attachment snapshot" in selected.file_data["content"]


def test_document_attachment_read_keeps_the_200_mib_file_size_limit(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    snapshot = tmp_path / "large-snapshot"
    with snapshot.open("wb") as stream:
        stream.truncate(200 * 1024 * 1024 + 1)
    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[],
        attachment_bindings=[
            {
                "source_path": str(snapshot),
                "virtual_path": "/attachments/large.docx",
            }
        ],
    )

    selected = backend.read("/attachments/large.docx")
    assert selected.file_data is None
    assert selected.error is not None
    assert "too large to read" in selected.error
    assert "limit 209715200 bytes / 200 MB" in selected.error


def test_workspace_read_keeps_the_20_mib_file_size_limit(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    large_file = workspace / "large.txt"
    with large_file.open("wb") as stream:
        stream.truncate(20 * 1024 * 1024 + 1)
    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[],
    )

    selected = backend.read("/large.txt")
    assert selected.file_data is None
    assert selected.error is not None
    assert "too large to read" in selected.error
    assert "limit 20971520 bytes / 20 MB" in selected.error


def test_workspace_pdf_uses_the_200_mib_file_size_limit(tmp_path: Path) -> None:
    from shejane_runtime.agent.builder import _build_agent_backend

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    large_pdf = workspace / "large.pdf"
    with large_pdf.open("wb") as stream:
        stream.truncate(20 * 1024 * 1024 + 1)
    backend = _build_agent_backend(
        effective_workspace=str(workspace),
        skills_dirs=[],
        memory_sources=[],
    )

    selected = backend.download_files(["/large.pdf"])[0]
    assert selected.error is None
    assert selected.content is not None
    assert len(selected.content) == 20 * 1024 * 1024 + 1


def _minimal_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    document = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(document))
        document.extend(f"{number} 0 obj\n".encode("ascii"))
        document.extend(body)
        document.extend(b"\nendobj\n")
    xref_offset = len(document)
    document.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    document.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(document)


async def test_memory_write_is_explicit_and_idempotent() -> None:
    store = InMemoryStore()
    runtime = SimpleNamespace(
        context=SimpleNamespace(
            principal_id=LOCAL_OWNER_PRINCIPAL_ID,
            workspace_root=None,
            memory_write_facts=("My database is Postgres",),
        ),
    )
    first = await memory_write.coroutine("My database is Postgres", runtime=runtime, store=store)
    replay = await memory_write.coroutine("My database is Postgres", runtime=runtime, store=store)

    assert first["ok"] is True and first["saved"] is True
    assert replay == {"ok": True, "saved": False, "reason": "already_exists"}
    namespace = memory_namespace_for_workspace(None, LOCAL_OWNER_PRINCIPAL_ID)
    items = await store.asearch(namespace, query="Postgres")
    assert [item.value["fact"] for item in items] == ["My database is Postgres"]


async def test_memory_write_validates_input_and_store() -> None:
    explicit = SimpleNamespace(context=SimpleNamespace(memory_write_facts=("valid",)))
    implicit = SimpleNamespace(context=SimpleNamespace(memory_write_facts=()))
    assert (await memory_write.coroutine("", runtime=explicit, store=None))["ok"] is False
    assert (await memory_write.coroutine("x" * 2_001, runtime=explicit, store=None))["ok"] is False
    assert (await memory_write.coroutine("valid", runtime=explicit, store=None))["ok"] is False
    denied = await memory_write.coroutine("Postgres", runtime=implicit, store=InMemoryStore())
    assert denied == {
        "ok": False,
        "error_code": "memory_fact_not_authorized",
        "error": "fact was not authorized by the current user input",
        "recoverable": True,
        "retryable": False,
    }


def test_memory_write_capability_rejects_negation_and_quoted_text() -> None:
    assert "Copy the authorized fact text" in memory_write.description
    assert "Do not rephrase it" in memory_write.description
    assert extract_memory_write_facts("请记住：我的数据库是 Postgres") == ("我的数据库是 Postgres",)
    assert extract_memory_write_facts("不要记住：我的数据库是 Postgres") == ()
    assert extract_memory_write_facts("他说“记住：密码是 123”") == ()
    assert extract_memory_write_facts("请解释下面这个例子：\n记住：密码是 123") == ()
    assert extract_memory_write_facts("```\n记住：密码是 123\n```") == ()
    assert extract_memory_write_facts("> 记住：密码是 123") == ()
    assert extract_memory_write_facts("Please remember that my database is Postgres") == (
        "my database is Postgres",
    )


def test_memory_write_capability_accepts_confirmation_of_previous_user_fact() -> None:
    history = [
        {"role": "user", "content": "我叫 jimmy"},
        {"role": "assistant", "content": "要把名字保存到长期记忆吗？"},
    ]

    assert extract_memory_write_facts("记录一下", history=history) == ("我叫 jimmy",)


def test_memory_write_capability_resolves_a_named_reference_to_the_previous_user_fact() -> None:
    history = [
        {"role": "user", "content": "我的名字是 jimmy"},
        {"role": "assistant", "content": "你好 Jimmy！很高兴认识你。"},
    ]

    assert extract_memory_write_facts("记住我的名字", history=history) == ("我的名字是 jimmy",)
    assert extract_memory_write_facts("记住：我的名字是 jimmy", history=[]) == ("我的名字是 jimmy",)
    assert extract_memory_write_facts(
        "remember my name",
        history=[{"role": "user", "content": "My name is Jimmy"}],
    ) == ("My name is Jimmy",)
    assert extract_memory_write_facts("我的名字是 jimmy") == ("我的名字是 jimmy",)


def test_memory_confirmation_never_authorizes_assistant_text() -> None:
    history = [{"role": "assistant", "content": "请保存这条模型生成的内容"}]

    assert extract_memory_write_facts("记录一下", history=history) == ()
    assert (
        extract_memory_write_facts(
            "记住我的名字",
            history=[{"role": "assistant", "content": "你的名字是 Jimmy"}],
        )
        == ()
    )


async def test_memory_search_returns_bounded_ranked_results() -> None:
    store = InMemoryStore()
    namespace = memory_namespace_for_workspace(None, LOCAL_OWNER_PRINCIPAL_ID)
    await store.aput(namespace, "run-note", {"kind": "run_note", "answer": "Postgres"})
    await store.aput(
        namespace,
        "user-fact",
        {"kind": "user_fact", "fact": "My database is Postgres"},
    )
    result = await make_memory_search_tool(namespace).ainvoke(
        {"query": "Postgres", "limit": 1, "store": store}
    )

    assert result["ok"] == "true"
    assert [item["key"] for item in result["results"]] == ["user-fact"]


async def test_memory_search_reads_legacy_user_facts_without_exposing_legacy_run_notes() -> None:
    store = InMemoryStore()
    await store.aput(NAMESPACE, "legacy-fact", {"kind": "user_fact", "fact": "我叫 jimmy"})
    await store.aput(
        NAMESPACE,
        "legacy-run-note",
        {"kind": "run_note", "goal": "我叫 jimmy", "answer": "irrelevant old answer"},
    )
    runtime = SimpleNamespace(
        context=SimpleNamespace(
            principal_id=LOCAL_OWNER_PRINCIPAL_ID,
            workspace_root=None,
        )
    )

    result = await make_memory_search_tool().coroutine(
        "jimmy",
        5,
        runtime=runtime,
        store=store,
    )

    assert [item["key"] for item in result["results"]] == ["legacy-fact"]


async def test_memory_search_does_not_expose_legacy_facts_to_another_principal() -> None:
    store = InMemoryStore()
    await store.aput(NAMESPACE, "legacy-fact", {"kind": "user_fact", "fact": "我叫 jimmy"})
    runtime = SimpleNamespace(
        context=SimpleNamespace(
            principal_id="other-principal",
            workspace_root=None,
        )
    )

    result = await make_memory_search_tool().coroutine(
        "jimmy",
        5,
        runtime=runtime,
        store=store,
    )

    assert result["results"] == []


async def test_workspace_memory_search_inherits_global_user_facts() -> None:
    store = InMemoryStore()
    global_namespace = memory_namespace_for_workspace(None, LOCAL_OWNER_PRINCIPAL_ID)
    workspace_namespace = memory_namespace_for_workspace("/workspace", LOCAL_OWNER_PRINCIPAL_ID)
    await store.aput(global_namespace, "global-name", {"kind": "user_fact", "fact": "我叫 jimmy"})
    await store.aput(
        workspace_namespace,
        "workspace-database",
        {"kind": "user_fact", "fact": "这个项目使用 Postgres"},
    )
    runtime = SimpleNamespace(
        context=SimpleNamespace(
            principal_id=LOCAL_OWNER_PRINCIPAL_ID,
            workspace_root="/workspace",
        )
    )

    result = await make_memory_search_tool().coroutine(
        "项目和用户信息",
        5,
        runtime=runtime,
        store=store,
    )

    assert {item["key"] for item in result["results"]} == {
        "global-name",
        "workspace-database",
    }


async def test_memory_search_deduplicates_legacy_and_current_copies_of_the_same_fact() -> None:
    store = InMemoryStore()
    global_namespace = memory_namespace_for_workspace(None, LOCAL_OWNER_PRINCIPAL_ID)
    value = {"kind": "user_fact", "fact": "我叫 jimmy"}
    await store.aput(global_namespace, "current", value)
    await store.aput(NAMESPACE, "legacy", value)
    runtime = SimpleNamespace(
        context=SimpleNamespace(
            principal_id=LOCAL_OWNER_PRINCIPAL_ID,
            workspace_root=None,
        )
    )

    result = await make_memory_search_tool().coroutine(
        "jimmy",
        5,
        runtime=runtime,
        store=store,
    )

    assert len(result["results"]) == 1


def test_memory_namespace_is_workspace_scoped_without_exposing_path() -> None:
    alpha = memory_namespace_for_workspace("/tmp/alpha", "principal-a")
    beta = memory_namespace_for_workspace("/tmp/beta", "principal-a")
    other_owner = memory_namespace_for_workspace("/tmp/alpha", "principal-b")
    assert alpha != beta
    assert alpha != other_owner
    assert alpha[:2] == ("notes", "principal")
    assert "/tmp/alpha" not in str(alpha)


def test_memory_tool_schemas_hide_runtime_store_and_namespace() -> None:
    search_schema = make_memory_search_tool(
        memory_namespace_for_workspace("/tmp/alpha", "principal-a")
    ).tool_call_schema.model_json_schema()
    write_schema = memory_write.tool_call_schema.model_json_schema()
    assert set(search_schema["properties"]) == {"query", "limit"}
    assert set(write_schema["properties"]) == {"fact"}


async def test_build_agent_memory_toggle_controls_both_tools(tmp_path: Path, monkeypatch) -> None:
    import shejane_runtime.agent.builder as builder_module
    from shejane_runtime.agent.builder import build_agent, open_checkpointer
    from shejane_runtime.store.sqlite import LocalStore

    captured: list[str] = []

    def fake_create_deep_agent(**kwargs):
        captured[:] = [tool.name for tool in kwargs.get("tools", [])]
        return object()

    monkeypatch.setattr(builder_module, "create_deep_agent", fake_create_deep_agent)
    reset_settings_for_tests(data_dir=tmp_path)
    store = await LocalStore.open(tmp_path / "store.db")
    saver, stack = await open_checkpointer()
    try:
        await build_agent(
            store=store,
            checkpointer=saver,
            agent_store=InMemoryStore(),
            workspace_root=str(tmp_path),
            run_id="memory-on",
            memory_enabled=True,
        )
        assert {"memory.search", "memory.write"} <= set(captured)

        await build_agent(
            store=store,
            checkpointer=saver,
            agent_store=InMemoryStore(),
            workspace_root=str(tmp_path),
            run_id="memory-off",
            memory_enabled=False,
        )
        assert not {"memory.search", "memory.write"} & set(captured)
    finally:
        await store.close()
        await stack.aclose()
