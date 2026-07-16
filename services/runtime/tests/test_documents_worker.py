from __future__ import annotations

import asyncio
import hashlib
import json
import os
import subprocess
import sys
import uuid
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.text.run import Run
from PIL import Image

from local_host.plugins.executor import ManagedWorkerActionExecutor
from local_host.plugins.macos_vm import load_macos_vm_resources
from local_host.plugins.platforms import (
    current_managed_worker_execution_platform,
    current_managed_worker_platform,
)
from local_host.plugins.runtime_assets import RuntimeAssetHandle

REPO_ROOT = Path(__file__).resolve().parents[3]
WORKER = REPO_ROOT / "plugins" / "office" / "documents" / "worker" / "documents_worker.py"
GOLDEN_ROOT = REPO_ROOT / "plugins" / "office" / "documents" / "golden"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _worker_command() -> tuple[str, ...]:
    frozen = os.environ.get("SHEJANE_TEST_DOCUMENTS_WORKER")
    return (frozen,) if frozen else (sys.executable, str(WORKER))


def _invocation(
    action_id: str,
    arguments: dict[str, object],
    *,
    source: Path | None = None,
) -> dict[str, object]:
    invocation: dict[str, object] = {
        "schema_version": 1,
        "invocation_id": str(uuid.uuid4()),
        "operation_id": f"run_01:{action_id}:001",
        "action": {
            "plugin_id": "org.shejane.documents",
            "plugin_digest": "sha256:" + "d" * 64,
            "action_id": action_id,
        },
        "arguments": arguments,
        "inputs": [],
        "grants": {"capabilities": []},
        "limits": {"timeout_ms": 10_000, "memory_mb": 512, "output_mb": 32},
    }
    if source is not None:
        data = source.read_bytes()
        invocation["inputs"] = [
            {
                "id": "document",
                "path": f"/input/{source.name}",
                "media_type": DOCX_MIME,
                "size_bytes": len(data),
                "sha256": hashlib.sha256(data).hexdigest(),
            }
        ]
    return invocation


async def _invoke(
    tmp_path: Path,
    action_id: str,
    arguments: dict[str, object],
    *,
    source: Path | None = None,
    runtime_assets: tuple[RuntimeAssetHandle, ...] = (),
) -> tuple[dict[str, object], Path]:
    roots = tmp_path / str(uuid.uuid4())
    input_root = roots / "input"
    output_root = roots / "output"
    input_root.mkdir(parents=True)
    output_root.mkdir()
    materialized = None
    if source is not None:
        materialized = input_root / source.name
        materialized.write_bytes(source.read_bytes())
    vm_manifest = os.environ.get("SHEJANE_TEST_MACOS_VM_ASSETS")
    command = _worker_command()
    result = await ManagedWorkerActionExecutor(
        command,
        runtime_assets=runtime_assets,
        vm_resources=(
            load_macos_vm_resources(Path(vm_manifest).resolve(strict=True)) if vm_manifest else None
        ),
        package_root=Path(command[0]).resolve(strict=True).parent if vm_manifest else None,
    ).invoke(
        _invocation(action_id, arguments, source=materialized),
        input_root=input_root,
        output_root=output_root,
    )
    return result, output_root


@pytest.mark.asyncio
async def test_documents_worker_create_read_and_edit_without_mutating_input(tmp_path: Path) -> None:
    create_arguments = {
        "filename": "created.docx",
        "metadata": {"title": "Quarterly note", "author": "SheJane"},
        "header": "Private header",
        "footer": "Page footer",
        "blocks": [
            {"type": "heading", "level": 1, "text": "Summary"},
            {"type": "paragraph", "text": "Revenue was stable."},
            {"type": "bullet_list", "items": ["North", "South"]},
            {"type": "table", "rows": [["Region", "Value"], ["North", "42"]]},
        ],
    }
    created, create_output = await _invoke(
        tmp_path,
        "document.create",
        create_arguments,
    )
    replay, replay_output = await _invoke(tmp_path, "document.create", create_arguments)
    created_path = create_output / "created.docx"
    assert created["status"] == "succeeded"
    assert replay["status"] == "succeeded"
    assert created_path.is_file()
    assert created_path.read_bytes() == (replay_output / "created.docx").read_bytes()

    document = Document(str(created_path))
    comment_run = document.paragraphs[1].runs[0]
    document.add_comment(comment_run, text="Check source", author="Reviewer")
    document.save(str(created_path))
    original_digest = hashlib.sha256(created_path.read_bytes()).hexdigest()

    read, _ = await _invoke(
        tmp_path,
        "document.read",
        {"input_id": "document", "include_markdown": True},
        source=created_path,
    )
    assert read["status"] == "succeeded"
    assert read["output"] == {
        "title": "Quarterly note",
        "author": "SheJane",
        "paragraph_count": 4,
        "table_count": 1,
        "comment_count": 1,
        "headings": [{"level": 1, "text": "Summary"}],
        "headers": ["Private header"],
        "footers": ["Page footer"],
        "markdown": (
            "# Summary\n\nRevenue was stable.\n\n- North\n\n- South\n\n"
            "| Region | Value |\n| --- | --- |\n| North | 42 |"
        ),
        "truncated": False,
        "warnings": [],
    }

    edited, edit_output = await _invoke(
        tmp_path,
        "document.edit",
        {
            "input_id": "document",
            "output_filename": "edited.docx",
            "operations": [
                {"type": "find_replace", "find": "stable", "replace": "growing"},
                {"type": "insert_paragraph", "after": "Revenue was growing.", "text": "Approved."},
                {"type": "set_footer", "text": "Final"},
                {"type": "set_metadata", "property": "subject", "value": "Board review"},
            ],
        },
        source=created_path,
    )
    edited_document = Document(str(edit_output / "edited.docx"))
    assert edited["status"] == "succeeded"
    assert [paragraph.text for paragraph in edited_document.paragraphs][:3] == [
        "Summary",
        "Revenue was growing.",
        "Approved.",
    ]
    assert edited_document.sections[0].footer.paragraphs[0].text == "Final"
    assert edited_document.core_properties.subject == "Board review"
    assert hashlib.sha256(created_path.read_bytes()).hexdigest() == original_digest


@pytest.mark.asyncio
async def test_documents_worker_returns_stable_input_errors(tmp_path: Path) -> None:
    for name, data, code in (
        ("broken.docx", b"not a zip", "document_corrupt"),
        ("legacy.docx", bytes.fromhex("d0cf11e0a1b11ae1"), "document_encrypted_or_legacy"),
    ):
        source = tmp_path / name
        source.write_bytes(data)
        result, _ = await _invoke(
            tmp_path,
            "document.read",
            {"input_id": "document"},
            source=source,
        )
        assert result["status"] == "failed"
        assert result["error"]["code"] == code  # type: ignore[index]


@pytest.mark.asyncio
async def test_documents_worker_renders_only_with_exact_runtime_asset_tools(tmp_path: Path) -> None:
    source = tmp_path / "render.docx"
    document = Document()
    document.add_paragraph("Rendered")
    document.save(str(source))

    asset_root = tmp_path / "office-runtime"
    asset_root.mkdir()
    soffice = asset_root / "soffice"
    mutool = asset_root / "mutool"
    soffice.write_text(
        f"#!{sys.executable}\n"
        "import pathlib, sys\n"
        "args = sys.argv[1:]\n"
        "out = pathlib.Path(args[args.index('--outdir') + 1])\n"
        "source = pathlib.Path(args[-1])\n"
        "(out / (source.stem + '.pdf')).write_bytes(b'%PDF-1.7\\n%%EOF\\n')\n",
        encoding="utf-8",
    )
    mutool.write_text(
        f"#!{sys.executable}\n"
        "import pathlib, sys\n"
        "args = sys.argv[1:]\n"
        "if args[0] == 'info':\n"
        "    print('Pages: 2')\n"
        "else:\n"
        "    pattern = args[args.index('-o') + 1]\n"
        "    for page in (1, 2):\n"
        "        pathlib.Path(pattern % page).write_bytes(b'PNG')\n",
        encoding="utf-8",
    )
    soffice.chmod(0o500)
    mutool.chmod(0o500)
    (asset_root / "office-runtime.json").write_text(
        json.dumps({"soffice": "soffice", "mutool": "mutool"}),
        encoding="utf-8",
    )
    sbom = asset_root / "sbom.spdx.json"
    sbom.write_text("{}", encoding="utf-8")
    asset = RuntimeAssetHandle(
        asset_id="org.libreoffice.runtime",
        version="25.8.7",
        platform="darwin/arm64",
        digest="sha256:" + "a" * 64,
        root=asset_root,
        payload=asset_root,
        license="MPL-2.0",
        source_url="https://www.libreoffice.org/",
        sbom=sbom,
    )

    rendered, output = await _invoke(
        tmp_path,
        "document.render",
        {"input_id": "document", "dpi": 144, "include_png": True, "max_pages": 2},
        source=source,
        runtime_assets=(asset,),
    )

    assert rendered["status"] == "succeeded", rendered.get("error", {}).get("message", rendered)
    assert rendered["output"] == {
        "page_count": 2,
        "rendered_pages": 2,
        "pdf_name": "render.pdf",
        "png_names": ["render.page-0001.png", "render.page-0002.png"],
        "warnings": [],
    }
    assert {path.name for path in output.iterdir()} == {
        "render.pdf",
        "render.page-0001.png",
        "render.page-0002.png",
    }


@pytest.mark.asyncio
async def test_documents_worker_real_runtime_asset_render(tmp_path: Path) -> None:
    configured = os.environ.get("SHEJANE_TEST_OFFICE_RUNTIME")
    if not configured:
        pytest.skip("real Office runtime asset not configured")
    asset_root = Path(configured).resolve(strict=True)
    target = (
        current_managed_worker_execution_platform()
        if os.environ.get("SHEJANE_TEST_MACOS_VM_ASSETS")
        else current_managed_worker_platform()
    )
    assert target is not None
    source = tmp_path / "cjk.docx"
    document = Document()
    document.core_properties.title = "SheJane Office golden"
    document.sections[0].header.paragraphs[0].text = "SheJane · 文档质量基线"
    document.sections[0].footer.paragraphs[0].text = "Deterministic Office Runtime"
    run = document.add_heading("跨平台文档", level=1).runs[0]
    _set_cjk_font(run)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("石间文档测试 · 日本語 · 한국어")
    _set_cjk_font(run)
    document.add_paragraph("第一项", style="List Bullet")
    document.add_paragraph("第二项", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "区域"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "华东"
    table.cell(1, 1).text = "42"
    document.add_page_break()
    run = document.add_heading("Second page", level=1).runs[0]
    _set_cjk_font(run)
    document.add_paragraph("Page breaks, headers, footers, lists, and tables remain stable.")
    document.save(str(source))
    asset = RuntimeAssetHandle(
        asset_id="org.libreoffice.runtime",
        version="25.8.7",
        platform=target,
        digest="sha256:" + "a" * 64,
        root=asset_root,
        payload=asset_root / "payload",
        license="MPL-2.0 AND AGPL-3.0-only AND OFL-1.1",
        source_url="https://www.libreoffice.org/",
        sbom=asset_root / "payload" / "office-runtime.json",
    )

    rendered, output = await _invoke(
        tmp_path,
        "document.render",
        {"input_id": "document", "dpi": 144, "include_png": True, "max_pages": 4},
        source=source,
        runtime_assets=(asset,),
    )

    assert rendered["status"] == "succeeded", rendered.get("error", {}).get("message", rendered)
    assert rendered["output"]["page_count"] == 2  # type: ignore[index]
    pdf = output / "cjk.pdf"
    assert pdf.read_bytes().startswith(b"%PDF")
    if not os.environ.get("SHEJANE_TEST_MACOS_VM_ASSETS"):
        config = json.loads((asset.payload / "office-runtime.json").read_text(encoding="utf-8"))
        mutool = asset.payload / config["mutool"]
        extracted = (
            await asyncio.to_thread(
                subprocess.run,
                [str(mutool), "draw", "-F", "txt", "-o", "-", str(pdf)],
                check=True,
                capture_output=True,
                text=True,
            )
        ).stdout
        assert "石间文档测试" in extracted
        assert "Second page" in extracted

    expected = json.loads((GOLDEN_ROOT / f"{target.replace('/', '-')}.json").read_text())
    actual = []
    for page in range(1, 3):
        with Image.open(output / f"cjk.page-{page:04d}.png") as image:
            gray = image.convert("L")
            actual.append(
                {
                    "size": list(gray.size),
                    "ink_bbox": list(_ink_bbox(gray)),
                    "dhash": _dhash(gray),
                }
            )
    assert [page["size"] for page in actual] == [page["size"] for page in expected["pages"]]
    for observed, baseline in zip(actual, expected["pages"], strict=True):
        assert all(
            abs(left - right) <= 2
            for left, right in zip(observed["ink_bbox"], baseline["ink_bbox"], strict=True)
        )
        assert _hamming(observed["dhash"], baseline["dhash"]) <= 8
    for observed, baseline in zip(actual, expected["pages"], strict=True):
        assert all(
            abs(left - right) <= 2
            for left, right in zip(observed["ink_bbox"], baseline["ink_bbox"], strict=True)
        )
        assert _hamming(observed["dhash"], baseline["dhash"]) <= 8


def _set_cjk_font(run: Run) -> None:
    run.font.name = "Noto Sans CJK SC"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")


def _ink_bbox(image: Image.Image) -> tuple[int, int, int, int]:
    inverted = image.point(lambda value: 255 - value)
    return inverted.getbbox() or (0, 0, 0, 0)


def _dhash(image: Image.Image) -> str:
    sampled = image.resize((9, 8), Image.Resampling.LANCZOS)
    pixels = list(sampled.get_flattened_data())
    value = 0
    for row in range(8):
        for column in range(8):
            value = (value << 1) | (pixels[row * 9 + column] > pixels[row * 9 + column + 1])
    return f"{value:016x}"


def _hamming(left: str, right: str) -> int:
    return (int(left, 16) ^ int(right, 16)).bit_count()
